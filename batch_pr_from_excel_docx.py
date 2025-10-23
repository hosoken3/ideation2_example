# -*- coding: utf-8 -*-
"""
batch_pr_from_excel_docx_interactive.py

Excelと参考PDFフォルダのパスは直書き。
実行後にCMDで行番号（例: 1,5,9）を入力 → 指定行ごとにWord（.docx）を出力。

要件:
  pip install openai pandas pdfplumber PyPDF2 openpyxl python-docx
  環境変数 OPENAI_API_KEY を設定
"""

import datetime
import json
from pathlib import Path
from typing import Dict, List, Optional

import pandas as pd
from docx import Document

# ====== 固定パス（必要に応じて書き換えてください） ======
EXCEL_PATH = r"C:\Users\cm000\OneDrive\ドキュメント\OpenAI\patent\src\4社統合\generated_ideas_with_process_20250924_2222.xlsx"
OUTDIR     = r"C:\Users\cm000\OneDrive\ドキュメント\OpenAI\patent\src\4社統合\out_pr_word"
REFS_DIR   = r"C:\Users\cm000\OneDrive\ドキュメント\OpenAI\patent\src\4社統合\プレスリリース作成"
MODEL      = "o3-2025-04-16"
PER_PDF_LIMIT = 6000     # 1ファイルあたり取り込み文字数上限
TOTAL_STYLE_LIMIT = 12000 # 参考テキスト合計上限
# ========================================================

REQUIRED_COLS = ["入力_技術ニーズ", "入力_製品シーズ", "アイデア名", "具体的なソリューション", "アイデア創出のプロセス"]

def read_pdf_text(pdf_path: Path) -> str:
    """pdfplumber優先、失敗時はPyPDF2でテキスト抽出。"""
    text = ""
    try:
        import pdfplumber
        with pdfplumber.open(str(pdf_path)) as pdf:
            for page in pdf.pages:
                txt = page.extract_text() or ""
                text += txt + "\n"
        if text.strip():
            return text
    except Exception:
        pass
    try:
        import PyPDF2
        with open(pdf_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for p in reader.pages:
                txt = p.extract_text() or ""
                text += txt + "\n"
    except Exception as e:
        print(f"⚠️ 参考PDFの読み取りに失敗: {pdf_path} / {e}")
    return text

def truncate(s: str, max_chars: int) -> str:
    s = s.replace("\r\n", "\n")
    while "\n\n\n" in s:
        s = s.replace("\n\n\n", "\n\n")
    return s[:max_chars]

def prepare_style_corpus(ref_dir: Path, per_pdf_limit: int = PER_PDF_LIMIT, total_limit: int = TOTAL_STYLE_LIMIT) -> Optional[str]:
    """フォルダ内のPDFを読み、抜粋を連結してスタイル参照テキストを作る。"""
    if not ref_dir.exists():
        print(f"⚠️ 参考PDFフォルダが見つかりません: {ref_dir}")
        return None
    snippets: List[str] = []
    for pdf_path in sorted(ref_dir.glob("*.pdf")):
        txt = read_pdf_text(pdf_path)
        if txt:
            snippets.append(truncate(txt, per_pdf_limit))
        else:
            print(f"⚠️ テキスト抽出できないPDFをスキップ: {pdf_path.name}")
    if not snippets:
        return None
    corpus = "\n\n---\n\n".join(snippets)
    return truncate(corpus, total_limit)

def sanitize(s: str) -> str:
    return (s or "").strip()

def build_fields_from_row(row: pd.Series) -> Dict[str, str]:
    needs_body = sanitize(row.get("入力_技術ニーズ", ""))
    seeds_body = sanitize(row.get("入力_製品シーズ", ""))
    idea_name  = sanitize(row.get("アイデア名", ""))
    solution   = sanitize(row.get("具体的なソリューション", ""))
    process    = sanitize(row.get("アイデア創出のプロセス", ""))
    return {
        "needs_news": "",
        "seeds_news": "",
        "needs_body": needs_body,
        "seeds_body": seeds_body,
        "idea": (idea_name + ("\n" + solution if solution else "")).strip(),
        "process": process,
    }

def call_openai(model: str, fields: Dict[str, str], date_jp: str, style_ref: Optional[str]) -> str:
    from openai import OpenAI
    client = OpenAI()

    system_prompt = (
        "あなたはローム株式会社の広報担当です。"
        "与えられた入力データと、参考リリース文の『構成・段落運び・トーン』を模倣しつつ、"
        "ただし文言のコピーペーストは避け、オリジナルの表現で、"
        "ローム公式にふさわしい事実基調で技術的信頼感のあるプレスリリースを作成してください。"
    )

    user_template = f"""
以下に与えられた技術情報（技術ニーズと製品シーズの組み合わせ）と、そこから導かれた「組み合わせアイデア」「技術的背景（プロセス）」をもとに、新製品のプレスリリース文を作成してください。

【目的】
ローム株式会社の公式プレスリリースの形式・構成・トーンに忠実に倣い、正確かつ客観的で技術的信頼感のある文章にしてください。

【分量】
全体で**1000〜1300字程度**を目安としてください。  
とくに「背景」や「特長」パートは**情報量・段落構成を充実させてください**。

【構成】
以下のフォーマットで出力してください：

---

### [タイトル]（製品名＋簡潔な価値）
〜 [キャッチコピー：1文、導入価値を象徴する表現] 〜  
発表日：{date_jp}

---

#### ＜要旨＞  
製品概要、開発目的、用途、製品名、発売計画を含めて1段落でまとめてください。

---

#### ＜背景＞  
社会・業界における技術課題やニーズ、市場動向を紹介し、従来技術の制限や課題も説明してください。

---

#### ＜製品の特長＞  
箇条書き形式で出力してください。各項目は**1段落以上の解説つき**でお願いします。

- 特長①：何が優れているか＋なぜそうなるか
- 特長②：どんな課題に対する解決か、他社・従来比の比較
- 特長③：設計者・エンジニアにとってのメリット
- 特長④：環境性能や持続可能性への寄与（ある場合）

---

#### ＜期待される効果＞  
製品がもたらすユーザーや社会への価値、導入効果、将来的な波及効果などを述べてください。

---

【入力データ】
技術ニーズ：
{fields.get('needs_body','')}

製品シーズ：
{fields.get('seeds_body','')}

組み合わせたアイデア：
{fields.get('idea','')}

プロセス（技術的背景）：
{fields.get('process','')}
""".strip()

    messages = [{"role": "system", "content": system_prompt}]
    if style_ref:
        messages.append({
            "role": "user",
            "content": (
                "以下は参考とするプレスリリースの抜粋です。語り口・段落構成・強調の仕方などを学習用の手がかりとして使用してください。"
                "ただし文言の直接引用は避けてください。\n\n"
                "【参考リリース（抜粋・複数PDF統合）】\n" + style_ref
            )
        })
    messages.append({"role": "user", "content": user_template})

    resp = client.chat.completions.create(model=model, messages=messages)
    return resp.choices[0].message.content.strip()

def mdish_to_docx(md_text: str, out_path: Path):
    """Markdown風の見出し・箇条書きを簡易変換して .docx へ出力。"""
    doc = Document()
    for line in md_text.splitlines():
        s = line.rstrip()
        if not s:
            doc.add_paragraph("")  # 空行
            continue
        if s.startswith("### "):
            doc.add_heading(s[4:].strip(), level=1)
        elif s.startswith("#### "):
            doc.add_heading(s[5:].strip(), level=2)
        elif s.startswith("- "):
            doc.add_paragraph(s[2:].strip(), style="List Bullet")
        else:
            # 区切り線(---)などは空行に
            if s.strip("-— ").strip() == "":
                doc.add_paragraph("")
            else:
                doc.add_paragraph(s)
    doc.save(str(out_path))

def main():
    print("=== プレスリリース生成（インタラクティブ） ===")
    excel_path = Path(EXCEL_PATH)
    outdir = Path(OUTDIR); outdir.mkdir(parents=True, exist_ok=True)

    # Excel読み込み & 列チェック
    df = pd.read_excel(excel_path, sheet_name=0)
    missing = [c for c in REQUIRED_COLS if c not in df.columns]
    if missing:
        raise SystemExit(f"Excelに必要列がありません: {missing}")

    # スタイル参照の準備（最新3〜5本のみをフォルダに入れておく運用を想定）
    style_text = prepare_style_corpus(Path(REFS_DIR))

    today = datetime.date.today()
    date_jp = f"{today.year}年{today.month:02d}月{today.day:02d}日（作成日）"

    # ユーザー入力（行番号）
    rows_input = input("生成したい行番号を入力してください（例: 1,5,9）：").strip()
    if not rows_input:
        print("⚠️ 行番号が入力されていません。終了します。")
        return
    rows = []
    for token in rows_input.split(","):
        token = token.strip()
        if token.isdigit():
            idx = int(token) - 1
            if 0 <= idx < len(df):
                rows.append(idx)
            else:
                print(f"⚠️ 範囲外の行番号をスキップ: {token}")
        else:
            print(f"⚠️ 数値でない入力をスキップ: {token}")
    if not rows:
        print("⚠️ 有効な行番号がありません。終了します。")
        return

    # 生成ループ
    for idx in rows:
        fields = build_fields_from_row(df.iloc[idx])
        content_md = call_openai(MODEL, fields, date_jp, style_text)
        out_path = outdir / f"press_release_row{idx+1:03d}.docx"
        mdish_to_docx(content_md, out_path)

        # 再現用データを保存（任意）
        with open(outdir / f"pr_input_row{idx+1:03d}.json", "w", encoding="utf-8") as jf:
            json.dump(fields, jf, ensure_ascii=False, indent=2)

        print(f"✅ 出力: {out_path}")

    print("=== 完了しました ===")

if __name__ == "__main__":
    main()
